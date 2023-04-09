from docx import Document
from docx.shared import Inches

# Create a new Word document
document = Document()

# Add a heading to the document
document.add_heading('PHP Basic', level=1)

# PHP Basic
for i in range(1, 103):
    filename = f"img/PHP Basic/ss/screenshot{i}.png"
    document.add_picture(filename, width=Inches(5.0))

document.add_heading('PHP Basic Algorithm', level=1)

# PHP Basic Algorithm
for i in range(1, 137):
    filename = f"img/PHP Basic Algorithm/ss/screenshot{i}.png"
    document.add_picture(filename, width=Inches(5.0))

document.add_heading('PHP arrays', level=1)

# PHP arrays
for i in range(1,60):
    filename = f"img/PHP arrays/ss/screenshot{i}.png"
    document.add_picture(filename, width=Inches(5.0))

document.add_heading('PHP for loop', level=1)

# PHP for loop
for i in range(1,39):
    filename = f"img/PHP for loop/ss/screenshot{i}.png"
    document.add_picture(filename, width=Inches(5.0))

document.add_heading('PHP functions', level=1)

# PHP functions
for i in range(1,7):
    filename = f"img/PHP functions/ss/screenshot{i}.png"
    document.add_picture(filename, width=Inches(5.0))

document.add_heading('PHP classes', level=1)

# PHP classes
for i in range(1,8):
    filename = f"img/PHP classes/ss/screenshot{i}.png"
    document.add_picture(filename, width=Inches(5.0))

document.add_heading('PHP Regular Expression', level=1)

# PHP Regular Expression
for i in range(1,8):
    filename = f"img/PHP Regular Expression/ss/screenshot{i}.png"
    document.add_picture(filename, width=Inches(5.0))

document.add_heading('PHP Date', level=1)

# PHP Date
for i in range(1,29):
    filename = f"img/PHP Date/ss/screenshot{i}.png"
    document.add_picture(filename, width=Inches(5.0))

document.add_heading('PHP String', level=1)

# PHP String
for i in range(1,27):
    filename = f"img/PHP String/ss/screenshot{i}.png"
    document.add_picture(filename, width=Inches(5.0))

document.add_heading('PHP Math', level=1)

# PHP Math
for i in range(1,13):
    filename = f"img/PHP Math/ss/screenshot{i}.png"
    document.add_picture(filename, width=Inches(5.0))

document.add_heading('PHP JSON', level=1)

# PHP JSON
for i in range(1,5):
    filename = f"img/PHP JSON/ss/screenshot{i}.png"
    document.add_picture(filename, width=Inches(5.0))

document.add_heading('PHP Searching and Sorting Algorithm', level=1)

# PHP Searching and Sorting Algorithm
for i in range(1,5):
    filename = f"img/PHP Searching and Sorting Algorithm/ss/screenshot{i}.png"
    document.add_picture(filename, width=Inches(5.0))

# Save the document
document.save('mydocument.docx')
